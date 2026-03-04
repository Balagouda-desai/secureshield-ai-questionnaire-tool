from flask import Flask, render_template, request, redirect, session
from werkzeug.utils import secure_filename
from utils.parser import extract_questions_from_pdf
from utils.retriever import load_reference_documents, chunk_documents, create_embeddings, find_best_match
import sqlite3
import os
from docx import Document
from flask import send_file

app = Flask(__name__)
app.secret_key = "supersecretkey"  # change later

DATABASE = "database.db"

def get_db_connection():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn

@app.route("/")
def home():
    if "user_id" in session:
        return redirect("/dashboard")
    return redirect("/login")

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        conn = get_db_connection()
        conn.execute("INSERT INTO users (username, password) VALUES (?, ?)",
                     (username, password))
        conn.commit()
        conn.close()

        return redirect("/login")

    return render_template("register.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        conn = get_db_connection()
        user = conn.execute("SELECT * FROM users WHERE username=? AND password=?",
                            (username, password)).fetchone()
        conn.close()

        if user:
            session["user_id"] = user["id"]
            return redirect("/dashboard")

    return render_template("login.html")

@app.route("/dashboard")
def dashboard():
    if "user_id" not in session:
        return redirect("/login")
    return render_template("dashboard.html")

@app.route("/upload_questionnaire", methods=["POST"])
def upload_questionnaire():

    if "user_id" not in session:
        return redirect("/login")

    file = request.files["questionnaire"]

    if file:
        filename = secure_filename(file.filename)

        folder = "uploads/questionnaires"

        # create folder if it doesn't exist
        os.makedirs(folder, exist_ok=True)

        filepath = os.path.join(folder, filename)

        file.save(filepath)

    return redirect("/dashboard")

@app.route("/upload_reference", methods=["POST"])
def upload_reference():

    if "user_id" not in session:
        return redirect("/login")

    file = request.files["reference"]

    if file:
        filename = secure_filename(file.filename)

        folder = "uploads/references"

        os.makedirs(folder, exist_ok=True)

        filepath = os.path.join(folder, filename)

        file.save(filepath)

    return redirect("/dashboard")

@app.route("/generate_answers")
def generate_answers():

    # get questionnaire file
    questionnaire_folder = "uploads/questionnaires"

    files = os.listdir(questionnaire_folder)

    if not files:
        return "No questionnaire uploaded."

    pdf_path = os.path.join(questionnaire_folder, files[0])

    # extract questions
    questions = extract_questions_from_pdf(pdf_path)


    # load reference documents
    reference_folder = "uploads/references"

    documents = load_reference_documents(reference_folder)


    # chunk documents
    chunks = chunk_documents(documents)


    # create embeddings
    chunk_embeddings = create_embeddings(chunks)


    results = []

    # process each question
    for question in questions:

        best_chunk, confidence = find_best_match(question, chunks, chunk_embeddings)

        chunk_text = best_chunk["text"]
        citation = best_chunk["filename"]

        answer = chunk_text
        evidence = chunk_text.split(".")[0] + "."

        if confidence < 0.40:
            answer = "Not found in references."
            citation = "None"
            evidence = "No supporting evidence found."

        results.append({
            "question": question,
            "answer": answer,
            "citation": citation,
            "confidence": float(confidence),
            "evidence": evidence
        })
    
    session["results"] = results

    total_questions = len(results)

    answered = sum(1 for r in results if r["citation"] != "None")

    not_found = total_questions - answered

    return render_template(
    "results.html",
    results=results,
    total_questions=total_questions,
    answered=answered,
    not_found=not_found
)

@app.route("/export_answers", methods=["POST"])
def export_answers():

    results = session.get("results")

    if not results:
        return "No results available."

    doc = Document()

    doc.add_heading("Answered Security Questionnaire", level=1)

    for item in results:

        doc.add_paragraph(item["question"], style="List Number")

        doc.add_paragraph("Answer:")
        doc.add_paragraph(item["answer"])

        doc.add_paragraph(f"Citation: {item['citation']}")
        doc.add_paragraph(f"Evidence: {item['evidence']}")
        doc.add_paragraph(f"Confidence: {item['confidence']:.2f}")

        doc.add_paragraph("")

    file_path = "answered_questionnaire.docx"

    doc.save(file_path)

    return send_file(file_path, as_attachment=True)

@app.route("/logout")
def logout():
    session.clear()
    return redirect("/login")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)